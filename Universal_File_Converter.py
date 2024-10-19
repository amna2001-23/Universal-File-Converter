import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx2pdf import convert
from PIL import Image
import pytesseract
import io
import os

# Conversion functions
def convert_pdf_to_word(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    
    doc = Document()
    doc.add_paragraph(text)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def convert_word_to_pdf(word_file):
    temp_docx_path = 'temp.docx'
    with open(temp_docx_path, 'wb') as f:
        f.write(word_file.read())
    
    try:
        convert(temp_docx_path, "converted.pdf")
        with open("converted.pdf", "rb") as f:
            pdf_data = f.read()
        return pdf_data
    except Exception as e:
        st.error(f"Error converting Word to PDF: {e}")
        return None
    finally:
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        if os.path.exists("converted.pdf"):
            os.remove("converted.pdf")

def convert_images_to_word(images):
    # Initialize a new Word document
    doc = Document()
    
    for image in images:
        # Convert image to text using OCR
        text = pytesseract.image_to_string(Image.open(image))

        # Add the text to the Word document with formatting
        doc.add_paragraph(text)
    
    # Save the Word document to an in-memory file
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# Streamlit app
st.sidebar.title('Navigation')
page = st.sidebar.radio('Go to', ['Introduction', 'Converter', 'Image to Word'])

if page == 'Introduction':
    st.title('Welcome to the Universal File Converter App')
    # Image slider
    with st.expander("Preview Images"):
        image_list = [
            'p2w.jpg',
            'p2w1.jpg',
            'w2p.png',
            'p2w2.jpg'
        ]
        image_index = st.slider('Slider', 0, len(image_list) - 1)
        st.image(image_list[image_index])
    
    st.subheader("Introducing Universal File Converter: The Ultimate File Converter")
    st.write("""
    This application allows you to convert files between Word and PDF formats and images to Word format.
    
    **Effortless Conversion**: Universal File Converter is a versatile file converter for all your document conversion needs.
    
    You can convert:
    - PDFs to Word
    - Word documents to PDFs
    - Images to Word (OCR conversion)

    Ready to experience the power of Universal File Converter? Try it today!
    """)
    
elif page == 'Converter':
    st.title('File Converter')
    
    conversion_type = st.selectbox("Select Conversion Type", ["PDF to Word", "Word to PDF"])
    
    if conversion_type == "PDF to Word":
        pdf_file = st.file_uploader("Upload PDF file", type=['pdf'])
        if pdf_file is not None:
            word_data = convert_pdf_to_word(pdf_file)
            st.download_button(label="Download Word file", data=word_data, file_name="converted.docx")
    
    elif conversion_type == "Word to PDF":
        word_file = st.file_uploader("Upload Word file", type=['docx'])
        if word_file is not None:
            pdf_data = convert_word_to_pdf(word_file)
            if pdf_data:
                st.download_button(label="Download PDF file", data=pdf_data, file_name="converted.pdf")

elif page == 'Image to Word':
    st.title('Image to Word Converter')

    # Allow multiple image uploads
    images = st.file_uploader("Upload Images", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    
    if images:
        if st.button('Convert to Word'):
            word_data = convert_images_to_word(images)
            st.download_button(label="Download Word file", data=word_data, file_name="converted.docx")

#streamlit run "Universal File Converter.py"
